"""
export.py — POST /api/export/docx

Produces a Word (.docx) compliance report: one table, one row per individual
FAIL/WARN issue, with a cropped "focus" image of the flagged element/pipe
(blank if no crop was captured for that row) and the non-compliant clause
description.

The frontend already has the full evaluation result (from /api/evaluate) and
captures each row's crop itself, directly from the live Konva canvas at high
resolution — see EvaluationModal.tsx / issueCropRegions.ts. This endpoint is
deliberately dumb: it just assembles whatever rows + crops it's given into a
table. It does not re-run compliance checks or do any image processing, so
there's no risk of the row content drifting from what the officer saw on
screen, and nothing here is persisted server-side.
"""

from __future__ import annotations

import io
import json

from docx import Document
from docx.shared import Inches
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse

from app.config import settings

router = APIRouter()


@router.post("/export/docx")
async def export_docx(
    manifest_json: str = Form(...),
    crops: list[UploadFile] = File(default=[]),
) -> StreamingResponse:
    try:
        rows: list[dict] = json.loads(manifest_json)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=422, detail=f"Invalid JSON in manifest_json: {e}")
    if not isinstance(rows, list):
        raise HTTPException(status_code=422, detail="manifest_json must be a JSON array.")
    if len(rows) > settings.max_report_rows:
        raise HTTPException(
            status_code=413,
            detail=f"manifest_json has {len(rows)} rows, over the {settings.max_report_rows} limit.",
        )

    if len(crops) > settings.max_crops:
        raise HTTPException(
            status_code=413,
            detail=f"{len(crops)} crops supplied, over the {settings.max_crops} limit.",
        )

    # Read one at a time against a running total rather than reading all of
    # them up front. Every crop stays in memory for the life of the request, so
    # the cumulative size is what bounds memory — a per-file cap would still
    # allow max_crops x per_file. See the security risk assessment (held outside this repo), R-03.
    crop_bytes: list[bytes] = []
    total = 0
    for f in crops:
        data = await f.read()
        total += len(data)
        if total > settings.max_total_crop_bytes:
            raise HTTPException(
                status_code=413,
                detail=(
                    f"crops exceed the {settings.max_total_crop_bytes // (1024 * 1024)}MB "
                    "combined limit."
                ),
            )
        crop_bytes.append(data)

    document = Document()
    document.add_heading("Compliance Evaluation — Non-Compliant Items", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Location"
    header[1].text = "Non-Compliant Clause"

    for row in rows:
        if not isinstance(row, dict):
            continue
        crop_index = row.get("crop_index")
        table_row = table.add_row()
        if isinstance(crop_index, int) and 0 <= crop_index < len(crop_bytes):
            run = table_row.cells[0].paragraphs[0].add_run()
            run.add_picture(io.BytesIO(crop_bytes[crop_index]), width=Inches(2.5))
        status = row.get("status", "")
        check_title = row.get("check_title", "")
        check_id = row.get("check_id", "")
        text = row.get("text", "")
        table_row.cells[1].text = f"[{status}] {check_title} ({check_id}): {text}"

    if len(rows) == 0:
        table.add_row().cells[1].text = "No FAIL or WARN items — schematic passed all checks."

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=compliance_report.docx"},
    )
